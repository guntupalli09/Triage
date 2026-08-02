"""Tests for docx_export.py's redlined-document generation.

python-docx's paragraph.text/paragraph.runs deliberately don't see inside
w:ins/w:del wrappers (it wasn't built with revision tracking in mind), so
these tests read the raw OOXML directly via lxml/docx's own element tree —
the same thing a real consumer (Word, or any other OOXML-aware tool) reads.

Decisions are keyed by finding_key ("{rule_id}#{index in the findings
list}"), not bare rule_id — see review_workflow.py's module docstring for
why: the same rule can fire more than once in one real document.
"""

import io
import zipfile

import pytest
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

try:
    import mammoth
except ImportError:
    mammoth = None

from docx_export import build_redlined_docx
from review_workflow import finding_key as fk


def _finding(rule_id, start, end, exact_snippet, redline_text=None, title=None, rationale="why it matters"):
    return {
        "rule_id": rule_id,
        "title": title or rule_id,
        "rationale": rationale,
        "start_index": start,
        "end_index": end,
        "exact_snippet": exact_snippet,
        "redline": {"suggested_redline": redline_text, "issue": rule_id, "legal_rationale": rationale} if redline_text else None,
    }


def _open(docx_bytes):
    return Document(io.BytesIO(docx_bytes))


def _all_text(doc, tag):
    """Concatenates the text of every element with this qualified tag
    (e.g. 'w:t' or 'w:delText') anywhere in the document body."""
    out = []
    for p in doc.paragraphs:
        for el in p._p.iter(qn(tag)):
            out.append(el.text or "")
    return "".join(out)


def _elements(doc, tag):
    out = []
    for p in doc.paragraphs:
        out.extend(p._p.iter(qn(tag)))
    return out


def _part_xml(docx_bytes, part_name):
    zf = zipfile.ZipFile(io.BytesIO(docx_bytes))
    return zf.read(part_name).decode("utf-8")


def _part_exists(docx_bytes, part_name):
    zf = zipfile.ZipFile(io.BytesIO(docx_bytes))
    return part_name in zf.namelist()


class TestTrackedInsertDelete:
    def test_accepted_redline_is_a_real_ins_and_del(self):
        text = "The Vendor shall have unlimited liability under this Agreement."
        findings = [_finding("H_LOL_01", 11, 32, "shall have unlimited", "shall cap aggregate")]
        decisions = {fk(0, "H_LOL_01"): {"action": "accepted"}}
        docx_bytes, skipped = build_redlined_docx("test.docx", text, findings, decisions, author="Jane Attorney")
        assert skipped == []
        doc = _open(docx_bytes)
        assert "shall cap aggregate" in _all_text(doc, "w:t")
        assert "shall have unlimited" in _all_text(doc, "w:delText")
        # and NOT as plain visible w:t (it must be a real deletion, not just struck-through display text)
        assert "shall have unlimited" not in _all_text(doc, "w:t")

    def test_ins_and_del_elements_present(self):
        text = "Original clause text here."
        findings = [_finding("R1", 0, 8, "Original", "Revised")]
        decisions = {fk(0, "R1"): {"action": "accepted"}}
        docx_bytes, _ = build_redlined_docx("test.docx", text, findings, decisions)
        doc = _open(docx_bytes)
        assert len(_elements(doc, "w:ins")) >= 1
        assert len(_elements(doc, "w:del")) >= 1

    def test_track_changes_author_is_the_reviewing_attorney(self):
        text = "Original clause text here."
        findings = [_finding("R1", 0, 8, "Original", "Revised")]
        decisions = {fk(0, "R1"): {"action": "accepted"}}
        docx_bytes, _ = build_redlined_docx("test.docx", text, findings, decisions, author="Jane Attorney")
        doc = _open(docx_bytes)
        ins_els = _elements(doc, "w:ins")
        del_els = _elements(doc, "w:del")
        assert all(el.get(qn("w:author")) == "Jane Attorney" for el in ins_els)
        assert all(el.get(qn("w:author")) == "Jane Attorney" for el in del_els)

    def test_falls_back_to_a_generic_author_if_none_given(self):
        text = "Original clause text here."
        findings = [_finding("R1", 0, 8, "Original", "Revised")]
        decisions = {fk(0, "R1"): {"action": "accepted"}}
        docx_bytes, _ = build_redlined_docx("test.docx", text, findings, decisions)
        doc = _open(docx_bytes)
        assert _elements(doc, "w:ins")[0].get(qn("w:author"))

    def test_ins_del_ids_are_unique(self):
        text = "AAAA BBBB CCCC DDDD"
        findings = [_finding("R1", 0, 4, "AAAA", "first-new"), _finding("R2", 10, 14, "CCCC", "second-new")]
        decisions = {fk(0, "R1"): {"action": "accepted"}, fk(1, "R2"): {"action": "accepted"}}
        docx_bytes, _ = build_redlined_docx("test.docx", text, findings, decisions)
        doc = _open(docx_bytes)
        ids = [el.get(qn("w:id")) for el in _elements(doc, "w:ins") + _elements(doc, "w:del")]
        assert len(ids) == len(set(ids)), "revision ids must be unique or Word can't tell them apart"


class TestDuplicateRuleId:
    """The same rule can fire more than once in one real document — deciding
    on one occurrence must never affect another occurrence of the same
    rule_id. This is a real bug found live against a real contract, not a
    hypothetical."""

    def test_accepting_one_occurrence_does_not_apply_to_another(self):
        text = "AAAA BBBB CCCC DDDD"
        findings = [
            _finding("H_ATTFEE_01", 0, 4, "AAAA", "first-redline"),
            _finding("H_ATTFEE_01", 10, 14, "CCCC", "second-redline"),
        ]
        # only accept the FIRST occurrence
        decisions = {fk(0, "H_ATTFEE_01"): {"action": "accepted"}}
        docx_bytes, skipped = build_redlined_docx("test.docx", text, findings, decisions)
        assert skipped == []
        doc = _open(docx_bytes)
        assert "first-redline" in _all_text(doc, "w:t")
        assert "second-redline" not in _all_text(doc, "w:t")
        # the second occurrence's original text must be untouched, not deleted
        assert "CCCC" in _all_text(doc, "w:t")
        assert "CCCC" not in _all_text(doc, "w:delText")

    def test_rejecting_one_occurrence_does_not_reject_another(self):
        text = "AAAA BBBB CCCC DDDD"
        findings = [
            _finding("H_ATTFEE_01", 0, 4, "AAAA", "first-redline"),
            _finding("H_ATTFEE_01", 10, 14, "CCCC", "second-redline"),
        ]
        decisions = {
            fk(0, "H_ATTFEE_01"): {"action": "accepted"},
            fk(1, "H_ATTFEE_01"): {"action": "rejected", "reason": "already mutual"},
        }
        docx_bytes, skipped = build_redlined_docx("test.docx", text, findings, decisions)
        assert skipped == []
        doc = _open(docx_bytes)
        assert "first-redline" in _all_text(doc, "w:t")  # accepted one applied
        assert "second-redline" not in _all_text(doc, "w:t")  # rejected one NOT applied
        comments_xml = _part_xml(docx_bytes, "word/comments.xml")
        assert comments_xml.count("<w:comment ") == 2  # one rationale comment + one reject-note comment
        assert "already mutual" in comments_xml


class TestEditedRedlines:
    def test_edited_text_used_instead_of_original_redline(self):
        text = "Some clause here that needs work."
        findings = [_finding("R1", 5, 11, "clause", "engine's suggestion")]
        decisions = {fk(0, "R1"): {"action": "edited", "edited_text": "attorney's custom wording"}}
        docx_bytes, _ = build_redlined_docx("test.docx", text, findings, decisions)
        doc = _open(docx_bytes)
        assert "attorney's custom wording" in _all_text(doc, "w:t")
        assert "engine's suggestion" not in _all_text(doc, "w:t")

    def test_edit_is_noted_in_the_comment_not_the_document_body(self):
        text = "Some clause here."
        findings = [_finding("R1", 5, 11, "clause", "engine text", rationale="original rationale")]
        decisions = {fk(0, "R1"): {"action": "edited", "edited_text": "new text"}}
        docx_bytes, _ = build_redlined_docx("test.docx", text, findings, decisions)
        # the body should read like a normal human edit — no visible "[edited]" tag
        assert "[edited by reviewer]" not in _all_text(_open(docx_bytes), "w:t")
        comments_xml = _part_xml(docx_bytes, "word/comments.xml")
        assert "edited by the reviewer" in comments_xml


class TestComments:
    def test_comment_part_exists_when_there_are_accepted_redlines(self):
        text = "Original clause text here."
        findings = [_finding("R1", 0, 8, "Original", "Revised", rationale="uncapped exposure")]
        decisions = {fk(0, "R1"): {"action": "accepted"}}
        docx_bytes, _ = build_redlined_docx("test.docx", text, findings, decisions)
        assert _part_exists(docx_bytes, "word/comments.xml")

    def test_comment_part_absent_when_nothing_applied(self):
        docx_bytes, _ = build_redlined_docx("test.docx", "plain text, no findings", [], {})
        assert not _part_exists(docx_bytes, "word/comments.xml")

    def test_comment_cites_the_rule_id_and_rationale(self):
        text = "Original clause text here."
        findings = [_finding("H_LOL_01", 0, 8, "Original", "Revised", rationale="uncapped exposure is risky")]
        decisions = {fk(0, "H_LOL_01"): {"action": "accepted"}}
        docx_bytes, _ = build_redlined_docx("test.docx", text, findings, decisions)
        comments_xml = _part_xml(docx_bytes, "word/comments.xml")
        assert "H_LOL_01" in comments_xml
        assert "uncapped exposure is risky" in comments_xml

    def test_comment_authored_by_the_engine_not_the_attorney(self):
        text = "Original clause text here."
        findings = [_finding("R1", 0, 8, "Original", "Revised")]
        decisions = {fk(0, "R1"): {"action": "accepted"}}
        docx_bytes, _ = build_redlined_docx("test.docx", text, findings, decisions, author="Jane Attorney")
        comments_xml = _part_xml(docx_bytes, "word/comments.xml")
        assert 'w:author="TriageCounsel Deterministic Engine"' in comments_xml
        assert 'w:author="Jane Attorney"' not in comments_xml  # the edit is Jane's; the rationale note is the engine's

    def test_content_types_and_relationship_are_registered(self):
        text = "Original clause text here."
        findings = [_finding("R1", 0, 8, "Original", "Revised")]
        decisions = {fk(0, "R1"): {"action": "accepted"}}
        docx_bytes, _ = build_redlined_docx("test.docx", text, findings, decisions)
        content_types = _part_xml(docx_bytes, "[Content_Types].xml")
        rels = _part_xml(docx_bytes, "word/_rels/document.xml.rels")
        assert "comments+xml" in content_types
        assert "relationships/comments" in rels

    def test_comment_reference_style_is_registered(self):
        text = "Original clause text here."
        findings = [_finding("R1", 0, 8, "Original", "Revised")]
        decisions = {fk(0, "R1"): {"action": "accepted"}}
        docx_bytes, _ = build_redlined_docx("test.docx", text, findings, decisions)
        styles = _part_xml(docx_bytes, "word/styles.xml")
        assert 'w:styleId="CommentReference"' in styles

    def test_one_comment_per_applied_redline(self):
        text = "AAAA BBBB CCCC DDDD"
        findings = [_finding("R1", 0, 4, "AAAA", "first-new"), _finding("R2", 10, 14, "CCCC", "second-new")]
        decisions = {fk(0, "R1"): {"action": "accepted"}, fk(1, "R2"): {"action": "accepted"}}
        docx_bytes, _ = build_redlined_docx("test.docx", text, findings, decisions)
        comments_xml = _part_xml(docx_bytes, "word/comments.xml")
        assert comments_xml.count("<w:comment ") == 2


class TestTrackChangesEnabled:
    def test_track_changes_turned_on_when_redlines_applied(self):
        text = "Original clause text here."
        findings = [_finding("R1", 0, 8, "Original", "Revised")]
        decisions = {fk(0, "R1"): {"action": "accepted"}}
        docx_bytes, _ = build_redlined_docx("test.docx", text, findings, decisions)
        settings = _part_xml(docx_bytes, "word/settings.xml")
        assert "<w:trackChanges" in settings

    def test_track_changes_turned_on_even_with_no_redlines(self):
        # harmless and still useful — an empty review still benefits from
        # having Track Changes on for whatever happens next
        docx_bytes, _ = build_redlined_docx("test.docx", "plain text", [], {})
        settings = _part_xml(docx_bytes, "word/settings.xml")
        assert "<w:trackChanges" in settings


class TestRejectedFindings:
    def test_rejected_redline_text_is_not_applied(self):
        text = "The Vendor shall have unlimited liability."
        findings = [_finding("R1", 11, 32, "shall have unlimited", "shall cap liability")]
        decisions = {fk(0, "R1"): {"action": "rejected", "reason": "not applicable"}}
        docx_bytes, skipped = build_redlined_docx("test.docx", text, findings, decisions)
        doc = _open(docx_bytes)
        assert "shall cap liability" not in _all_text(doc, "w:t")
        assert "shall have unlimited" in _all_text(doc, "w:t")  # untouched plain text, not a deletion
        assert skipped == []
        # original text is untouched — no deletion was made, even though a comment exists
        assert "shall have unlimited" not in _all_text(doc, "w:delText")

    def test_rejected_finding_gets_a_comment_with_the_reason(self):
        text = "The Vendor shall have unlimited liability under this Agreement."
        findings = [_finding("H_LOL_01", 11, 32, "shall have unlimited", "shall cap liability", title="Unlimited Liability")]
        decisions = {fk(0, "H_LOL_01"): {"action": "rejected", "reason": "deal size too small to matter"}}
        docx_bytes, skipped = build_redlined_docx("test.docx", text, findings, decisions)
        assert _part_exists(docx_bytes, "word/comments.xml")
        comments_xml = _part_xml(docx_bytes, "word/comments.xml")
        assert "H_LOL_01" in comments_xml
        assert "deal size too small to matter" in comments_xml
        assert "Reviewed and declined" in comments_xml

    def test_rejected_finding_comment_is_authored_by_the_attorney_not_the_engine(self):
        text = "The Vendor shall have unlimited liability."
        findings = [_finding("R1", 11, 32, "shall have unlimited", "shall cap liability")]
        decisions = {fk(0, "R1"): {"action": "rejected", "reason": "not applicable"}}
        docx_bytes, _ = build_redlined_docx("test.docx", text, findings, decisions, author="Jane Attorney")
        comments_xml = _part_xml(docx_bytes, "word/comments.xml")
        assert 'w:author="Jane Attorney"' in comments_xml
        assert 'w:author="TriageCounsel Deterministic Engine"' not in comments_xml

    def test_rejected_finding_produces_no_ins_or_del(self):
        text = "The Vendor shall have unlimited liability."
        findings = [_finding("R1", 11, 32, "shall have unlimited", "shall cap liability")]
        decisions = {fk(0, "R1"): {"action": "rejected", "reason": "not applicable"}}
        docx_bytes, _ = build_redlined_docx("test.docx", text, findings, decisions)
        doc = _open(docx_bytes)
        assert _elements(doc, "w:ins") == []
        assert _elements(doc, "w:del") == []

    def test_rejected_and_accepted_findings_both_annotated_in_one_pass(self):
        text = "AAAA BBBB CCCC DDDD"
        findings = [
            _finding("R1", 0, 4, "AAAA", "first-new", title="First"),
            _finding("R2", 10, 14, "CCCC", None, title="Second"),
        ]
        decisions = {fk(0, "R1"): {"action": "accepted"}, fk(1, "R2"): {"action": "rejected", "reason": "no thanks"}}
        docx_bytes, skipped = build_redlined_docx("test.docx", text, findings, decisions)
        assert skipped == []
        doc = _open(docx_bytes)
        assert "first-new" in _all_text(doc, "w:t")
        assert "CCCC" in _all_text(doc, "w:t")  # rejected span untouched
        comments_xml = _part_xml(docx_bytes, "word/comments.xml")
        assert comments_xml.count("<w:comment ") == 2
        assert "no thanks" in comments_xml

    def test_rejected_with_empty_reason_is_not_annotated(self):
        # shouldn't happen given review_workflow's validation, but must not crash or
        # produce a blank comment if it somehow does
        text = "The Vendor shall have unlimited liability."
        findings = [_finding("R1", 11, 32, "shall have unlimited", "shall cap liability")]
        decisions = {fk(0, "R1"): {"action": "rejected", "reason": "  "}}
        docx_bytes, skipped = build_redlined_docx("test.docx", text, findings, decisions)
        assert not _part_exists(docx_bytes, "word/comments.xml")

    def test_rejected_comment_initials_reflect_the_attorney_not_the_engine(self):
        text = "The Vendor shall have unlimited liability."
        findings = [_finding("R1", 11, 32, "shall have unlimited", "shall cap liability")]
        decisions = {fk(0, "R1"): {"action": "rejected", "reason": "not applicable"}}
        docx_bytes, _ = build_redlined_docx("test.docx", text, findings, decisions, author="Jane Attorney")
        comments_xml = _part_xml(docx_bytes, "word/comments.xml")
        assert 'w:initials="JA"' in comments_xml

    def test_rejected_finding_overlapping_an_accepted_one_is_skipped(self):
        text = "0123456789ABCDEFGHIJ"
        findings = [
            _finding("R1", 0, 10, text[0:10], "FIRST-REPLACEMENT"),
            _finding("R2", 5, 15, text[5:15], None),
        ]
        decisions = {fk(0, "R1"): {"action": "accepted"}, fk(1, "R2"): {"action": "rejected", "reason": "overlaps"}}
        docx_bytes, skipped = build_redlined_docx("test.docx", text, findings, decisions)
        assert skipped == ["R2"]


class TestOtherSkippedDecisions:
    def test_flagged_findings_with_no_redline_are_not_applied(self):
        text = "Some indemnification clause text."
        findings = [_finding("R1", 5, 22, "indemnification")]
        decisions = {fk(0, "R1"): {"action": "flagged"}}
        docx_bytes, skipped = build_redlined_docx("test.docx", text, findings, decisions)
        doc = _open(docx_bytes)
        assert "indemnification" in _all_text(doc, "w:t")

    def test_undecided_finding_is_untouched(self):
        text = "Text with a finding nobody decided on yet."
        findings = [_finding("R1", 5, 9, "with", "changed")]
        docx_bytes, skipped = build_redlined_docx("test.docx", text, findings, {})
        doc = _open(docx_bytes)
        assert "changed" not in _all_text(doc, "w:t")
        assert skipped == []

    def test_missing_position_data_is_skipped_not_crashed(self):
        findings = [{"rule_id": "R1", "start_index": None, "end_index": None, "exact_snippet": "x", "redline": {"suggested_redline": "y"}}]
        decisions = {fk(0, "R1"): {"action": "accepted"}}
        docx_bytes, skipped = build_redlined_docx("test.docx", "some text", findings, decisions)
        assert docx_bytes


class TestOverlappingSpans:
    def test_overlapping_accepted_findings_the_second_is_skipped_not_corrupted(self):
        text = "0123456789ABCDEFGHIJ"
        findings = [
            _finding("R1", 0, 10, text[0:10], "FIRST-REPLACEMENT"),
            _finding("R2", 5, 15, text[5:15], "SECOND-REPLACEMENT"),
        ]
        decisions = {fk(0, "R1"): {"action": "accepted"}, fk(1, "R2"): {"action": "accepted"}}
        docx_bytes, skipped = build_redlined_docx("test.docx", text, findings, decisions)
        assert skipped == ["R2"]
        doc = _open(docx_bytes)
        assert "FIRST-REPLACEMENT" in _all_text(doc, "w:t")
        assert "SECOND-REPLACEMENT" not in _all_text(doc, "w:t")

    def test_non_overlapping_findings_both_applied_in_order(self):
        text = "AAAA BBBB CCCC DDDD"
        findings = [_finding("R1", 0, 4, "AAAA", "first-new"), _finding("R2", 10, 14, "CCCC", "second-new")]
        decisions = {fk(0, "R1"): {"action": "accepted"}, fk(1, "R2"): {"action": "accepted"}}
        docx_bytes, skipped = build_redlined_docx("test.docx", text, findings, decisions)
        assert skipped == []
        full_text = _all_text(_open(docx_bytes), "w:t")
        assert "first-new" in full_text and "second-new" in full_text
        assert full_text.index("first-new") < full_text.index("second-new")


class TestDocumentStructure:
    def test_header_note_present(self):
        docx_bytes, _ = build_redlined_docx("test.docx", "text", [], {})
        full_text = _all_text(_open(docx_bytes), "w:t")
        assert "Track Changes" in full_text

    def test_filename_in_title(self):
        docx_bytes, _ = build_redlined_docx("High.docx", "text", [], {})
        assert "High.docx" in _all_text(_open(docx_bytes), "w:t")

    def test_paragraph_breaks_preserved(self):
        text = "First paragraph.\n\nSecond paragraph."
        docx_bytes, _ = build_redlined_docx("test.docx", text, [], {})
        texts = [p.text for p in _open(docx_bytes).paragraphs if p.text.strip()]
        assert any(t == "First paragraph." for t in texts)
        assert any(t == "Second paragraph." for t in texts)

    def test_empty_document_does_not_crash(self):
        docx_bytes, skipped = build_redlined_docx("test.docx", "", [], {})
        assert docx_bytes
        assert skipped == []


class TestPackageIntegrity:
    def test_output_is_a_well_formed_zip(self):
        text = "Original clause text here."
        findings = [_finding("R1", 0, 8, "Original", "Revised")]
        decisions = {fk(0, "R1"): {"action": "accepted"}}
        docx_bytes, _ = build_redlined_docx("test.docx", text, findings, decisions)
        zf = zipfile.ZipFile(io.BytesIO(docx_bytes))
        assert zf.testzip() is None  # None means every member's CRC checked out

    def test_every_injected_part_is_well_formed_xml(self):
        text = "Original clause text here."
        findings = [_finding("R1", 0, 8, "Original", "Revised")]
        decisions = {fk(0, "R1"): {"action": "accepted"}}
        docx_bytes, _ = build_redlined_docx("test.docx", text, findings, decisions)
        zf = zipfile.ZipFile(io.BytesIO(docx_bytes))
        for name in ("[Content_Types].xml", "word/_rels/document.xml.rels", "word/styles.xml",
                     "word/document.xml", "word/settings.xml", "word/comments.xml"):
            etree.fromstring(zf.read(name))  # raises on malformed XML

    def test_document_reopens_cleanly_in_python_docx(self):
        text = "Original clause text here. More text follows."
        findings = [_finding("R1", 0, 8, "Original", "Revised")]
        decisions = {fk(0, "R1"): {"action": "accepted"}}
        docx_bytes, _ = build_redlined_docx("test.docx", text, findings, decisions)
        doc = _open(docx_bytes)  # must not raise
        assert len(doc.paragraphs) > 0


class TestAgainstRealEngineOutput:
    def test_real_finding_positions_round_trip(self, rule_engine):
        text = "In no event shall liability be limited under this Agreement, without limit."
        result = rule_engine.analyze(text)
        findings = [
            {
                "rule_id": f.rule_id, "title": f.title, "rationale": f.rationale,
                "start_index": f.start_index, "end_index": f.end_index,
                "exact_snippet": f.exact_snippet, "redline": None,
            }
            for f in result["findings"]
        ]
        if not findings:
            pytest.skip("no findings on this fixture text")
        decisions = {fk(0, findings[0]["rule_id"]): {"action": "flagged"}}
        docx_bytes, skipped = build_redlined_docx("test.docx", text, findings, decisions)
        assert docx_bytes


@pytest.mark.skipif(mammoth is None, reason="mammoth not installed — pip install -r requirements-dev.txt")
class TestIllegalXmlCharacterSanitization:
    """Regression coverage for the full-corpus-benchmark-discovered bug:
    contract text containing characters illegal in XML 1.0 (NUL bytes, C0
    control characters — see docx_export._ILLEGAL_XML_CHARS_RE) crashed
    build_redlined_docx with a raw lxml ValueError instead of producing a
    valid document. Every case here must, at minimum:
      1. Export without raising.
      2. Reopen cleanly in python-docx.
      3. Parse cleanly with mammoth (an independent, unrelated OOXML parser).
      4. Have every OOXML part be well-formed XML.
      5. Preserve Track Changes (w:ins/w:del) and comment anchoring exactly
         as the equivalent clean-text case would.
      6. Read back with the illegal character(s) removed and every other
         character — including legitimate "exotic" ones: smart quotes, em
         dashes, tabs, CR, supplementary-plane Unicode — unchanged.
    """

    def _validate_and_open(self, docx_bytes: bytes):
        """Shared validation every case below runs: well-formed zip, every
        OOXML part well-formed XML, python-docx reopen, mammoth parse with
        zero errors. Returns the reopened python-docx Document."""
        zf = zipfile.ZipFile(io.BytesIO(docx_bytes))
        assert zf.testzip() is None
        for name in zf.namelist():
            if name.endswith(".xml") or name.endswith(".rels"):
                etree.fromstring(zf.read(name))  # raises on malformed XML
        doc = _open(docx_bytes)
        mammoth_result = mammoth.convert_to_html(io.BytesIO(docx_bytes))
        assert not mammoth_result.messages, f"mammoth reported issues: {mammoth_result.messages}"
        return doc

    def test_embedded_null_byte_in_untracked_text(self):
        """The exact shape of the original benchmark failure: a NUL byte in
        plain (non-redlined) contract text flowing through _flush_text."""
        text = "Recital clause with a stray byte: \x00 right in the middle. End of recital."
        docx_bytes, skipped = build_redlined_docx("ex-21.htm", text, [], {})
        doc = self._validate_and_open(docx_bytes)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "\x00" not in full_text
        assert "Recital clause with a stray byte:" in full_text
        assert "right in the middle. End of recital." in full_text

    def test_all_illegal_c0_control_characters_stripped(self):
        """Every C0 control character XML 1.0 actually forbids (tab/LF/CR
        are legal and deliberately excluded from this set)."""
        illegal_codepoints = list(range(0x00, 0x09)) + [0x0B, 0x0C] + list(range(0x0E, 0x20))
        junk = "".join(chr(cp) for cp in illegal_codepoints)
        text = f"Before{junk}After"
        docx_bytes, _ = build_redlined_docx("test.docx", text, [], {})
        doc = self._validate_and_open(docx_bytes)
        full_text = "".join(p.text for p in doc.paragraphs)
        for cp in illegal_codepoints:
            assert chr(cp) not in full_text, f"illegal codepoint {hex(cp)} survived into the document"
        assert "Before" in full_text and "After" in full_text

    def test_mixed_legal_and_illegal_characters_preserves_legal_ones(self):
        text = "Party A\x00 shall\x01 pay\x02 Party B\x03 within\x04 30\x05 days\x06."
        docx_bytes, _ = build_redlined_docx("test.docx", text, [], {})
        doc = self._validate_and_open(docx_bytes)
        full_text = "".join(p.text for p in doc.paragraphs)
        assert "Party A shall pay Party B within 30 days." in full_text
        for cp in range(0x00, 0x07):
            assert chr(cp) not in full_text

    def test_unicode_edge_cases_bmp_noncharacters_stripped_supplementary_plane_kept(self):
        # U+FFFE / U+FFFD-adjacent noncharacters are illegal XML Char values;
        # a supplementary-plane character (an emoji) is fully legal and must
        # survive untouched.
        text = "Fee￾ structure \U0001F4B0 applies to all invoices."
        docx_bytes, _ = build_redlined_docx("test.docx", text, [], {})
        doc = self._validate_and_open(docx_bytes)
        full_text = "".join(p.text for p in doc.paragraphs)
        assert "￾" not in full_text
        assert "\U0001F4B0" in full_text
        assert "Fee" in full_text and "structure" in full_text and "applies to all invoices." in full_text

    def test_smart_quotes_and_em_dash_are_legal_and_preserved(self):
        text = "The Vendor’s obligations — as defined herein — include “prompt” delivery."
        docx_bytes, _ = build_redlined_docx("test.docx", text, [], {})
        doc = self._validate_and_open(docx_bytes)
        full_text = "".join(p.text for p in doc.paragraphs)
        assert "’" in full_text  # right single quote
        assert "—" in full_text  # em dash
        assert "“" in full_text and "”" in full_text  # curly double quotes

    def test_tabs_preserved(self):
        text = "Column A\tColumn B\tColumn C values here."
        docx_bytes, _ = build_redlined_docx("test.docx", text, [], {})
        doc = self._validate_and_open(docx_bytes)
        full_text = "".join(p.text for p in doc.paragraphs)
        assert "\t" in full_text

    def test_carriage_returns_do_not_crash_and_content_survives(self):
        text = "Line one.\r\nLine two.\r\nLine three with a null\x00 in it."
        docx_bytes, _ = build_redlined_docx("test.docx", text, [], {})
        doc = self._validate_and_open(docx_bytes)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "\x00" not in full_text
        assert "Line one." in full_text and "Line two." in full_text and "Line three with a null" in full_text

    def test_multi_line_clause_with_illegal_characters_across_paragraphs(self):
        text = (
            "SECTION 1. DEFINITIONS\n\n"
            "\"Confidential Information\" means any\x00 data disclosed\x01 by either party.\n\n"
            "SECTION 2. TERM\n\n"
            "This Agreement is effective as of the date first written above\x02."
        )
        docx_bytes, _ = build_redlined_docx("test.docx", text, [], {})
        doc = self._validate_and_open(docx_bytes)
        assert len(doc.paragraphs) >= 4
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "\x00" not in full_text and "\x01" not in full_text and "\x02" not in full_text
        assert "SECTION 1. DEFINITIONS" in full_text
        assert "SECTION 2. TERM" in full_text

    def test_very_large_paragraph_with_scattered_illegal_characters(self):
        # ~50k characters, an illegal byte scattered every ~500 chars —
        # exercises the same code path as the real 845KB/128K-word contract
        # that originally triggered this bug in the full-corpus benchmark.
        chunk = "This is a representative sentence of a large commercial contract clause. "
        illegal = "\x00\x0B\x1F"
        pieces = []
        for i in range(700):
            pieces.append(chunk)
            if i % 10 == 0:
                pieces.append(illegal[i % len(illegal)])
        text = "".join(pieces)
        assert len(text) > 50_000
        docx_bytes, _ = build_redlined_docx("large.docx", text, [], {})
        doc = self._validate_and_open(docx_bytes)
        full_text = "".join(p.text for p in doc.paragraphs)
        assert "\x00" not in full_text and "\x0B" not in full_text and "\x1F" not in full_text
        assert full_text.count("This is a representative sentence") == 700

    def test_overlapping_redlines_containing_illegal_characters(self):
        """Two findings whose spans overlap AND whose underlying text
        contains illegal characters: the overlap-skip logic and the
        sanitization must both apply correctly — the surviving (first,
        left-to-right) redline's tracked change must be clean, and the
        skipped one must not corrupt the document."""
        text = "Liability\x00 is uncapped\x01 and unlimited\x02 for both parties in all cases."
        findings = [
            _finding("R1", 0, 26, "Liability\x00 is uncapped\x01", "Liability is capped", rationale="rationale one\x00"),
            _finding("R2", 10, 40, "is uncapped\x01 and unlimited\x02", "is capped", rationale="rationale two\x01"),
        ]
        decisions = {
            fk(0, "R1"): {"action": "accepted"},
            fk(1, "R2"): {"action": "accepted"},
        }
        docx_bytes, skipped = build_redlined_docx("test.docx", text, findings, decisions)
        assert skipped == ["R2"]  # second span starts before the first ends -> skipped, unchanged from non-illegal-char behavior
        doc = self._validate_and_open(docx_bytes)

        ins_texts = _all_text(doc, "w:t")
        del_texts = _all_text(doc, "w:delText")
        assert "\x00" not in ins_texts and "\x01" not in ins_texts and "\x02" not in ins_texts
        assert "\x00" not in del_texts and "\x01" not in del_texts and "\x02" not in del_texts
        assert "Liability is capped" in ins_texts
        assert "Liability" in del_texts

        comments_xml = _part_xml(docx_bytes, "word/comments.xml")
        assert "\x00" not in comments_xml
        assert "rationale one" in comments_xml

    def test_accepted_reading_unchanged_except_for_illegal_xml_normalization(self):
        """The core correctness property: sanitization must be exactly
        equivalent to reading the same text with illegal characters already
        stripped — nothing else about content, structure, or track-changes
        behavior may differ."""
        from docx_export import _sanitize_xml_text

        dirty = "Payment\x00 due within\x0B thirty (30) days\x1F of invoice\x00 receipt."
        clean = _sanitize_xml_text(dirty)
        assert clean == "Payment due within thirty (30) days of invoice receipt."

        docx_dirty, _ = build_redlined_docx("test.docx", dirty, [], {})
        docx_clean, _ = build_redlined_docx("test.docx", clean, [], {})

        doc_dirty = self._validate_and_open(docx_dirty)
        doc_clean = _open(docx_clean)
        text_dirty = "".join(p.text for p in doc_dirty.paragraphs)
        text_clean = "".join(p.text for p in doc_clean.paragraphs)
        assert text_dirty == text_clean

    def test_illegal_characters_in_filename_heading_do_not_crash(self):
        docx_bytes, _ = build_redlined_docx("weird\x00name.htm", "Some contract text here.", [], {})
        doc = self._validate_and_open(docx_bytes)
        heading_text = doc.paragraphs[0].text
        assert "\x00" not in heading_text
        assert "weirdname.htm" in heading_text or "weird" in heading_text

    def test_track_changes_and_comment_attachment_intact_with_illegal_characters(self):
        """Confirms Track Changes + comment anchoring survive sanitization
        exactly as they do in the clean-text case — same element counts,
        same relationships, just without the illegal bytes."""
        text = "The Contractor\x00 shall indemnify\x01 the Company for all claims\x02 arising hereunder."
        findings = [_finding("R1", 0, 40, "The Contractor\x00 shall indemnify\x01", "The Contractor shall not indemnify", rationale="mutual indemnity\x00 expected")]
        decisions = {fk(0, "R1"): {"action": "accepted"}}
        docx_bytes, skipped = build_redlined_docx("test.docx", text, findings, decisions)
        assert skipped == []
        doc = self._validate_and_open(docx_bytes)

        ins_elements = _elements(doc, "w:ins")
        del_elements = _elements(doc, "w:del")
        assert len(ins_elements) == 1
        assert len(del_elements) == 1

        comment_refs = []
        for p in doc.paragraphs:
            comment_refs.extend(p._p.iter(qn("w:commentReference")))
        assert len(comment_refs) == 1

        comments_xml = _part_xml(docx_bytes, "word/comments.xml")
        assert comments_xml.count("<w:comment ") == 1
        assert "\x00" not in comments_xml and "\x01" not in comments_xml
