"""
Playbook Workbench — Phase 1 manual authoring routes.

A FastAPI APIRouter included into main.py's `app` (see
`app.include_router(playbook_workbench.router)` in main.py), kept in its
own module rather than growing main.py's already-large route list
further. Nothing here touches main.py's existing /playbooks/* routes,
PolicyRule, or any contract-review evaluation call site — see
docs/architecture/playbook_authoring_ux_design.md and
playbook_authoring.py's module docstring for the boundary this respects.
Production contract analysis continues reading PolicyRule exclusively;
PolicyPosition is authoring-only until Phase 4.

Every route re-derives the user from the session cookie and re-queries
the playbook filtered by that user's id — no route ever trusts a
playbook_id, clause_type, or position id from the URL/body as sufficient
authorization on its own. Ownership is re-checked on every request.
"""
from __future__ import annotations

import io
import os
from datetime import datetime
from typing import Any, Dict, Optional

from docx import Document
from fastapi import APIRouter, Depends, File, Form, HTTPException, Request, UploadFile
from fastapi.responses import HTMLResponse, RedirectResponse
from fastapi.templating import Jinja2Templates
from PyPDF2 import PdfReader
from sqlalchemy.orm import Session as DBSession

import audit_log
import google_oauth
import playbook_ai_extraction as pai
import playbook_authoring as pa
import playbook_extraction as pex
import policy_enforcement
import upload_security
from auth import get_current_user
from csrf import csrf_protect, get_csrf_token
from database import get_db
from models import Playbook, PlaybookSourceDocument, PolicyPosition, PolicyPositionApproval
from rules_engine import RuleEngine

_rule_engine = RuleEngine()

router = APIRouter()

# A second Jinja2Templates instance pointed at the same directory as
# main.py's — not shared module state, to avoid a circular import
# (main.py imports this module; this module cannot import `templates`
# back from main.py). Registers the same two globals main.py's instance
# does, since templates/base_app.html depends on both.
templates = Jinja2Templates(directory="templates")
templates.env.globals["google_signin_enabled"] = google_oauth.is_configured()
templates.env.globals["csrf_token"] = get_csrf_token


def _require_user(request: Request, db: DBSession):
    user = get_current_user(request, db)
    if not user:
        raise HTTPException(status_code=302, headers={"Location": "/login"})
    return user


def _get_owned_playbook(db: DBSession, user, playbook_id: int) -> Playbook:
    """The single ownership check every route in this module goes
    through — filtered by both id AND user_id in one query, so a
    nonexistent playbook and someone else's playbook are indistinguishable
    (both 404, never a 403 that would confirm the id exists)."""
    playbook = db.query(Playbook).filter(Playbook.id == playbook_id, Playbook.user_id == user.id).first()
    if not playbook:
        raise HTTPException(status_code=404, detail="Playbook not found")
    return playbook


def _require_clause_type(clause_type: str) -> str:
    if clause_type not in pa.CLAUSE_TYPES:
        raise HTTPException(status_code=404, detail="Unknown clause type")
    return clause_type


def _field_evidence_summary(position: Optional[PolicyPosition]) -> Dict[str, Dict[str, Any]]:
    """Per-field {source, status, excerpt, document_filename} for every
    current (non-superseded) field on a position — this is the answer to
    "why does TriageCounsel think this is our policy" (Phase 2 task item
    3): every EXTRACTED/CONFLICTING field carries its own evidence
    excerpt through to this dict, never just a bare value."""
    if position is None:
        return {}
    out: Dict[str, Dict[str, Any]] = {}
    for f in position.fields:
        if f.superseded_by_field_id is not None:
            continue
        doc_filename = None
        if f.evidence_document_id:
            doc = next((d for d in position.playbook.source_documents if d.id == f.evidence_document_id), None)
            doc_filename = doc.original_filename if doc else None
        out[f.field_name] = {
            "source": f.source, "status": f.status, "excerpt": f.evidence_excerpt,
            "document_filename": doc_filename, "extraction_version": f.extraction_version,
        }
    return out


def _base_context(request: Request, user, playbook: Playbook, clause_type: str, position: Optional[PolicyPosition]) -> dict:
    return {
        "request": request, "user": user, "playbook": playbook,
        "clause_type": clause_type, "clause_label": pa.CLAUSE_TYPE_LABELS[clause_type],
        "position": position, "cfg": (position.config_json or {}) if position else {},
        "field_statuses": pa.current_field_statuses(position) if position else {},
        "field_evidence": _field_evidence_summary(position),
        "field_labels": pa.FIELD_LABELS[clause_type],
        "shared_field_labels": pa.SHARED_FIELD_LABELS,
        "vocab": {f: pa.vocabulary_for(clause_type, f) for f in pa.CLAUSE_TYPE_CONFIG_FIELDS[clause_type]},
        "is_new_revision": False,
        "current_year": datetime.now().year,
        "error": None,
        # Shared-field values the form should re-display. Normally the
        # persisted position's; overridden with what was just submitted when
        # a save/submit fails validation, so nothing typed is lost (P0-1).
        "form_contract_side": position.contract_side if position else "mutual",
        "form_escalation_approval_authority": (position.escalation_approval_authority if position else None) or "",
        "form_fallback_text": (position.fallback_text if position else None) or "",
    }


def _apply_submitted_values(ctx: dict, clause_type: str, form) -> None:
    """Re-render an authoring form with the values the lawyer actually
    submitted rather than whatever is (or isn't) persisted — the whole
    point of P0-1's "if validation fails, preserve entered values"
    requirement."""
    ctx["cfg"] = pa.parse_clause_form_best_effort(clause_type, form)
    ctx["field_statuses"] = {name: "ESTABLISHED" for name in ctx["cfg"]}
    ctx["form_contract_side"] = pa.parse_contract_side(form.get("contract_side"))
    ctx["form_escalation_approval_authority"] = (form.get("escalation_approval_authority") or "").strip()
    ctx["form_fallback_text"] = (form.get("fallback_text") or "").strip()


def _extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """Duplicated from main.py's extract_text_from_file (not imported, to
    avoid a circular import — main.py imports this module to mount its
    router). Same behavior, same choke point (magic-byte validation,
    malware scan, zip/PDF-bomb guards via upload_security) as every other
    upload path in the app; kept in sync by hand since the two call sites
    serve different route trees."""
    ext = os.path.splitext(filename.lower())[1]

    upload_security.validate_magic_bytes(file_bytes, ext)
    upload_security.scan_for_malware(file_bytes)

    if ext == ".txt":
        try:
            return file_bytes.decode("utf-8", errors="ignore")
        except Exception:
            return file_bytes.decode("latin-1", errors="ignore")
    if ext == ".pdf":
        reader = PdfReader(io.BytesIO(file_bytes))
        upload_security.validate_pdf_page_count(len(reader.pages))
        text = "\n".join(p.extract_text() or "" for p in reader.pages)
        return upload_security.enforce_extracted_text_limit(text)
    if ext == ".docx":
        upload_security.validate_docx_zip_safety(file_bytes)
        doc = Document(io.BytesIO(file_bytes))
        text = "\n".join(p.text for p in doc.paragraphs)
        return upload_security.enforce_extracted_text_limit(text)
    raise ValueError("Unsupported file type")


# ---------------------------------------------------------------------------
# Workbench
# ---------------------------------------------------------------------------

@router.get("/playbooks/{playbook_id}/workbench", response_class=HTMLResponse)
async def playbook_workbench(request: Request, playbook_id: int, db: DBSession = Depends(get_db)):
    user = _require_user(request, db)
    playbook = _get_owned_playbook(db, user, playbook_id)
    coverage = pa.compute_coverage(db, playbook)
    cards = []
    for c in coverage.clauses:
        missing = []
        if c.position is not None and c.position.status in ("NEEDS_REVIEW", "APPROVED", "DRAFT"):
            try:
                pa.validate_position_for_activation(c.position)
            except pa.PolicyActivationError as exc:
                missing = pa.missing_field_labels(c.clause_type, exc.missing_fields)
        cards.append({
            "clause_type": c.clause_type, "label": c.label, "status_bucket": c.status_bucket,
            "position": c.position, "headline": pa.card_headline(c.position),
            "missing_required": missing,
        })
    return templates.TemplateResponse("playbook_workbench.html", {
        "request": request, "user": user, "playbook": playbook,
        "coverage": coverage, "cards": cards, "current_year": datetime.now().year,
        # Lifecycle status ("Active") is not production authority — see
        # policy_enforcement.is_policy_authoritative (P0-2).
        "enforcement": policy_enforcement.enforcement_disclosure(),
    })


# ---------------------------------------------------------------------------
# Authoring — edit / save
# ---------------------------------------------------------------------------

@router.get("/playbooks/{playbook_id}/positions/{clause_type}/edit", response_class=HTMLResponse)
async def position_edit_page(request: Request, playbook_id: int, clause_type: str, db: DBSession = Depends(get_db)):
    user = _require_user(request, db)
    playbook = _get_owned_playbook(db, user, playbook_id)
    clause_type = _require_clause_type(clause_type)

    # Idempotent: only actually creates a new row the first time this is
    # opened after the position reached ACTIVE (see
    # get_or_build_editable_position's docstring) — a page refresh never
    # forks a second revision.
    position, is_new_revision = pa.get_or_build_editable_position(db, playbook, clause_type)
    db.commit()

    ctx = _base_context(request, user, playbook, clause_type, position)
    ctx["is_new_revision"] = is_new_revision
    return templates.TemplateResponse(f"policy_position_fields/{clause_type}.html", ctx)


@router.post("/playbooks/{playbook_id}/positions/{clause_type}/save", response_class=HTMLResponse)
async def position_save(
    request: Request, playbook_id: int, clause_type: str,
    db: DBSession = Depends(get_db), _csrf: None = Depends(csrf_protect),
):
    user = _require_user(request, db)
    playbook = _get_owned_playbook(db, user, playbook_id)
    clause_type = _require_clause_type(clause_type)
    form = await request.form()

    position, is_new_revision = pa.get_or_build_editable_position(db, playbook, clause_type)

    try:
        clause_updates = pa.parse_clause_form(clause_type, form)
        pa.apply_position_update(
            db, position, clause_field_updates=clause_updates,
            contract_side=pa.parse_contract_side(form.get("contract_side")),
            escalation_approval_authority=(form.get("escalation_approval_authority") or "").strip() or None,
            fallback_text=(form.get("fallback_text") or "").strip() or None,
            user=user,
        )
    except (pa.PositionFormError, pa.PolicyConfigValidationError) as exc:
        db.rollback()
        # Preserve what was typed (P0-1) rather than re-rendering a blank
        # form — the values only exist in this request body.
        ctx = _base_context(request, user, playbook, clause_type, pa.get_position_for_display(db, playbook.id, clause_type))
        _apply_submitted_values(ctx, clause_type, form)
        ctx["is_new_revision"] = is_new_revision
        ctx["error"] = str(exc)
        return templates.TemplateResponse(f"policy_position_fields/{clause_type}.html", ctx, status_code=400)

    db.commit()
    audit_log.record_event(
        db, "policy_position_saved", request=request, actor_user_id=user.id,
        target_type="policy_position", target_id=position.id, success=True,
        metadata={"clause_type": clause_type, "playbook_id": playbook.id},
    )
    db.commit()
    return RedirectResponse(url=f"/playbooks/{playbook.id}/positions/{clause_type}/edit", status_code=302)


# ---------------------------------------------------------------------------
# Lifecycle transitions
# ---------------------------------------------------------------------------

def _get_current_position_or_404(db: DBSession, playbook: Playbook, clause_type: str) -> PolicyPosition:
    position = pa.get_position_for_display(db, playbook.id, clause_type)
    if position is None:
        raise HTTPException(status_code=404, detail="No policy position exists yet for this clause type")
    return position


@router.post("/playbooks/{playbook_id}/positions/{clause_type}/submit-for-review", response_class=HTMLResponse)
async def position_submit_for_review(
    request: Request, playbook_id: int, clause_type: str,
    db: DBSession = Depends(get_db), _csrf: None = Depends(csrf_protect),
):
    """Persist-then-transition, atomically, in one request (UX walkthrough
    P0-1). This route used to read ONLY the last-persisted row, so a lawyer
    who filled in the authoring form and clicked "Submit for review" without
    first clicking "Save draft" silently submitted the stale row and lost
    every value on screen.

    The authoring form now posts its full field payload here (via
    formaction, marked with the `authoring_form` hidden input), and this
    route runs the exact same persist step position_save runs — same
    pa.parse_clause_form / pa.apply_position_update pair, not a second
    parallel implementation — before transitioning. If validation fails,
    the transaction is rolled back, the form is re-rendered with the
    submitted values and an actionable error, and NO lifecycle transition
    occurs.

    A post that carries no authoring payload (no `authoring_form` marker —
    e.g. the review page's own transition action, or an API caller that
    already saved) keeps the previous transition-only behavior."""
    user = _require_user(request, db)
    playbook = _get_owned_playbook(db, user, playbook_id)
    clause_type = _require_clause_type(clause_type)

    form = await request.form()
    carries_form_state = bool(form.get("authoring_form"))

    if carries_form_state:
        position, _is_new_revision = pa.get_or_build_editable_position(db, playbook, clause_type)
    else:
        position = _get_current_position_or_404(db, playbook, clause_type)

    try:
        if carries_form_state:
            pa.apply_position_update(
                db, position,
                clause_field_updates=pa.parse_clause_form(clause_type, form),
                contract_side=pa.parse_contract_side(form.get("contract_side")),
                escalation_approval_authority=(form.get("escalation_approval_authority") or "").strip() or None,
                fallback_text=(form.get("fallback_text") or "").strip() or None,
                user=user,
            )
            db.flush()
        pa.mark_ready_for_review(db, position, user)
    except (pa.PositionFormError, pa.PositionLifecycleError, pa.PolicyConfigValidationError) as exc:
        db.rollback()
        # Re-render with what the lawyer actually typed, never a blank form
        # or a stale DB row — losing the input is the bug being fixed.
        ctx = _base_context(request, user, playbook, clause_type, pa.get_position_for_display(db, playbook.id, clause_type))
        if carries_form_state:
            _apply_submitted_values(ctx, clause_type, form)
        ctx["error"] = str(exc)
        return templates.TemplateResponse(f"policy_position_fields/{clause_type}.html", ctx, status_code=400)

    db.commit()
    audit_log.record_event(
        db, "policy_position_submitted_for_review", request=request, actor_user_id=user.id,
        target_type="policy_position", target_id=position.id, success=True,
        metadata={"clause_type": clause_type, "playbook_id": playbook.id},
    )
    db.commit()
    return RedirectResponse(url=f"/playbooks/{playbook.id}/positions/{clause_type}/review", status_code=302)


@router.post("/playbooks/{playbook_id}/positions/{clause_type}/return-to-draft", response_class=HTMLResponse)
async def position_return_to_draft(
    request: Request, playbook_id: int, clause_type: str,
    db: DBSession = Depends(get_db), _csrf: None = Depends(csrf_protect),
):
    user = _require_user(request, db)
    playbook = _get_owned_playbook(db, user, playbook_id)
    clause_type = _require_clause_type(clause_type)
    position = _get_current_position_or_404(db, playbook, clause_type)

    try:
        pa.return_to_draft(db, position, user)
    except pa.PositionLifecycleError as exc:
        raise HTTPException(status_code=400, detail=str(exc))

    db.commit()
    audit_log.record_event(
        db, "policy_position_returned_to_draft", request=request, actor_user_id=user.id,
        target_type="policy_position", target_id=position.id, success=True,
        metadata={"clause_type": clause_type, "playbook_id": playbook.id},
    )
    db.commit()
    return RedirectResponse(url=f"/playbooks/{playbook.id}/positions/{clause_type}/edit", status_code=302)


@router.get("/playbooks/{playbook_id}/positions/{clause_type}/review", response_class=HTMLResponse)
async def position_review_page(request: Request, playbook_id: int, clause_type: str, db: DBSession = Depends(get_db)):
    user = _require_user(request, db)
    playbook = _get_owned_playbook(db, user, playbook_id)
    clause_type = _require_clause_type(clause_type)
    position = _get_current_position_or_404(db, playbook, clause_type)

    missing_required = []
    try:
        pa.validate_position_for_activation(position)
    except pa.PolicyActivationError as exc:
        missing_required = pa.missing_field_labels(clause_type, exc.missing_fields)

    history = (
        db.query(PolicyPositionApproval)
        .filter(PolicyPositionApproval.policy_position_id == position.id)
        .order_by(PolicyPositionApproval.created_at.desc())
        .all()
    )

    ctx = _base_context(request, user, playbook, clause_type, position)
    ctx["summary_lines"] = pa.summarize_position(position)
    ctx["missing_required"] = missing_required
    ctx["history"] = history
    return templates.TemplateResponse("playbook_position_review.html", ctx)


@router.post("/playbooks/{playbook_id}/positions/{clause_type}/approve", response_class=HTMLResponse)
async def position_approve(
    request: Request, playbook_id: int, clause_type: str,
    db: DBSession = Depends(get_db), _csrf: None = Depends(csrf_protect),
):
    user = _require_user(request, db)
    playbook = _get_owned_playbook(db, user, playbook_id)
    clause_type = _require_clause_type(clause_type)
    position = _get_current_position_or_404(db, playbook, clause_type)

    try:
        pa.approve_position(db, position, user)
    except pa.PolicyActivationError as exc:
        db.rollback()
        ctx = _base_context(request, user, playbook, clause_type, position)
        ctx["summary_lines"] = pa.summarize_position(position)
        ctx["missing_required"] = pa.missing_field_labels(clause_type, exc.missing_fields)
        ctx["history"] = (
            db.query(PolicyPositionApproval)
            .filter(PolicyPositionApproval.policy_position_id == position.id)
            .order_by(PolicyPositionApproval.created_at.desc()).all()
        )
        ctx["error"] = "This position isn't ready to approve yet — see the unanswered questions below."
        return templates.TemplateResponse("playbook_position_review.html", ctx, status_code=400)
    except pa.PositionLifecycleError as exc:
        db.rollback()
        raise HTTPException(status_code=400, detail=str(exc))

    db.commit()
    audit_log.record_event(
        db, "policy_position_approved", request=request, actor_user_id=user.id,
        target_type="policy_position", target_id=position.id, success=True,
        metadata={"clause_type": clause_type, "playbook_id": playbook.id},
    )
    db.commit()
    return RedirectResponse(url=f"/playbooks/{playbook.id}/positions/{clause_type}/review", status_code=302)


@router.post("/playbooks/{playbook_id}/positions/{clause_type}/activate", response_class=HTMLResponse)
async def position_activate(
    request: Request, playbook_id: int, clause_type: str,
    db: DBSession = Depends(get_db), _csrf: None = Depends(csrf_protect),
):
    user = _require_user(request, db)
    playbook = _get_owned_playbook(db, user, playbook_id)
    clause_type = _require_clause_type(clause_type)
    position = _get_current_position_or_404(db, playbook, clause_type)

    try:
        pa.activate_position(db, position, user)
    except pa.PolicyActivationError as exc:
        db.rollback()
        ctx = _base_context(request, user, playbook, clause_type, position)
        ctx["summary_lines"] = pa.summarize_position(position)
        ctx["missing_required"] = pa.missing_field_labels(clause_type, exc.missing_fields)
        ctx["history"] = (
            db.query(PolicyPositionApproval)
            .filter(PolicyPositionApproval.policy_position_id == position.id)
            .order_by(PolicyPositionApproval.created_at.desc()).all()
        )
        ctx["error"] = "This position isn't ready to activate yet — see the unanswered questions below."
        return templates.TemplateResponse("playbook_position_review.html", ctx, status_code=400)
    except pa.PositionLifecycleError as exc:
        db.rollback()
        raise HTTPException(status_code=400, detail=str(exc))

    db.commit()
    audit_log.record_event(
        db, "policy_position_activated", request=request, actor_user_id=user.id,
        target_type="policy_position", target_id=position.id, success=True,
        metadata={"clause_type": clause_type, "playbook_id": playbook.id},
    )
    db.commit()
    return RedirectResponse(url=f"/playbooks/{playbook.id}/workbench", status_code=302)


# ---------------------------------------------------------------------------
# Preview — read-only, no Contract/review record, no production effect
# ---------------------------------------------------------------------------

@router.get("/playbooks/{playbook_id}/positions/{clause_type}/preview", response_class=HTMLResponse)
async def position_preview_page(request: Request, playbook_id: int, clause_type: str, db: DBSession = Depends(get_db)):
    user = _require_user(request, db)
    playbook = _get_owned_playbook(db, user, playbook_id)
    clause_type = _require_clause_type(clause_type)
    position = _get_current_position_or_404(db, playbook, clause_type)

    return templates.TemplateResponse("playbook_position_preview.html", {
        "request": request, "user": user, "playbook": playbook, "clause_type": clause_type,
        "clause_label": pa.CLAUSE_TYPE_LABELS[clause_type], "position": position,
        "sample_text": "", "decision": None, "current_year": datetime.now().year,
    })


@router.post("/playbooks/{playbook_id}/positions/{clause_type}/preview", response_class=HTMLResponse)
async def position_preview_run(
    request: Request, playbook_id: int, clause_type: str,
    db: DBSession = Depends(get_db), _csrf: None = Depends(csrf_protect),
):
    user = _require_user(request, db)
    playbook = _get_owned_playbook(db, user, playbook_id)
    clause_type = _require_clause_type(clause_type)
    position = _get_current_position_or_404(db, playbook, clause_type)

    form = await request.form()
    sample_text = (form.get("sample_text") or "").strip()
    decision = pa.run_preview(position, sample_text) if sample_text else None

    # Preview is read-only by construction: run_preview never calls
    # db.add/db.commit, and this route doesn't either. No Contract or
    # review record exists as a result of this request.
    return templates.TemplateResponse("playbook_position_preview.html", {
        "request": request, "user": user, "playbook": playbook, "clause_type": clause_type,
        "clause_label": pa.CLAUSE_TYPE_LABELS[clause_type], "position": position,
        "sample_text": sample_text, "decision": decision, "current_year": datetime.now().year,
    })


# ---------------------------------------------------------------------------
# Phase 2 — Deterministic/Private Template Import
# ---------------------------------------------------------------------------

ALLOWED_IMPORT_EXTENSIONS = {".txt", ".pdf", ".docx"}


@router.get("/playbooks/{playbook_id}/import", response_class=HTMLResponse)
async def playbook_import_page(request: Request, playbook_id: int, db: DBSession = Depends(get_db)):
    user = _require_user(request, db)
    playbook = _get_owned_playbook(db, user, playbook_id)
    return templates.TemplateResponse("playbook_import.html", {
        "request": request, "user": user, "playbook": playbook,
        "error": None, "current_year": datetime.now().year,
    })


@router.post("/playbooks/{playbook_id}/import", response_class=HTMLResponse)
async def playbook_import_submit(
    request: Request, playbook_id: int,
    file: UploadFile = File(...),
    use_as_deviation_baseline: str = Form(""),
    use_for_policy_extraction: str = Form(""),
    contract_side: str = Form("mutual"),
    db: DBSession = Depends(get_db), _csrf: None = Depends(csrf_protect),
):
    user = _require_user(request, db)
    playbook = _get_owned_playbook(db, user, playbook_id)
    contract_side = pa.parse_contract_side(contract_side)

    def _error(message: str):
        return templates.TemplateResponse("playbook_import.html", {
            "request": request, "user": user, "playbook": playbook,
            "error": message, "current_year": datetime.now().year,
        }, status_code=400)

    if not file.filename:
        return _error("Please choose a file to upload.")
    ext = os.path.splitext(file.filename.lower())[1]
    if ext not in ALLOWED_IMPORT_EXTENSIONS:
        return _error("Unsupported file type. Upload a .txt, .pdf, or .docx file.")
    if not use_as_deviation_baseline and not use_for_policy_extraction:
        return _error("Select at least one use for this document.")

    file_bytes = await file.read()
    try:
        extracted_text = _extract_text_from_file(file_bytes, file.filename)
    except Exception:
        return _error("Could not read this file. Make sure it's a valid .txt, .pdf, or .docx document.")

    source_document = PlaybookSourceDocument(
        playbook_id=playbook.id, uploaded_by_user_id=user.id, document_type="TEMPLATE_CONTRACT",
        original_filename=upload_security.sanitize_filename(file.filename), extracted_text=extracted_text,
        use_as_deviation_baseline=(use_as_deviation_baseline == "on"),
        use_for_policy_extraction=(use_for_policy_extraction == "on"),
    )
    db.add(source_document)
    db.flush()

    if source_document.use_as_deviation_baseline:
        # Same deterministic mechanism the legacy /playbooks/new upload
        # path already uses (rules_engine.analyze) — Phase 2 does not
        # introduce a second way to populate this, per design doc §8.2.
        analysis = _rule_engine.analyze(extracted_text)
        playbook.template_text = extracted_text
        playbook.template_findings_json = [
            {"rule_id": f.rule_id, "rule_name": f.rule_name, "title": f.title,
             "severity": f.severity.value, "rationale": f.rationale,
             "matched_excerpt": f.matched_excerpt}
            for f in analysis["findings"]
        ]
        playbook.template_risk = analysis["overall_risk"]

    touched_clause_types = []
    if source_document.use_for_policy_extraction:
        # Deterministic, private, regex-based extraction only — the same
        # extract_*_facts() functions the six engines already ship with.
        # No network call, no LLM, no document content leaves this
        # process. See playbook_extraction.import_source_document.
        positions = pex.import_source_document(db, playbook, source_document, contract_side, user)
        touched_clause_types = list(positions.keys())

    db.commit()
    audit_log.record_event(
        db, "playbook_source_document_imported", request=request, actor_user_id=user.id,
        target_type="playbook_source_document", target_id=source_document.id, success=True,
        metadata={
            "playbook_id": playbook.id, "use_as_deviation_baseline": source_document.use_as_deviation_baseline,
            "use_for_policy_extraction": source_document.use_for_policy_extraction,
            "clause_types_proposed": touched_clause_types,
        },
    )
    db.commit()

    if source_document.use_for_policy_extraction:
        return RedirectResponse(url=f"/playbooks/{playbook.id}/import/{source_document.id}/review", status_code=302)
    return RedirectResponse(url=f"/playbooks/{playbook.id}/workbench", status_code=302)


def _display_value(value: Any) -> str:
    if value is True:
        return "Yes"
    if value is False:
        return "No"
    if isinstance(value, list):
        return ", ".join(str(v).replace("_", " ") for v in value) if value else "None"
    if value is None:
        return ""
    return str(value)


def _get_owned_source_document(db: DBSession, playbook: Playbook, document_id: int) -> PlaybookSourceDocument:
    doc = db.query(PlaybookSourceDocument).filter(
        PlaybookSourceDocument.id == document_id, PlaybookSourceDocument.playbook_id == playbook.id,
    ).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Source document not found")
    return doc


@router.get("/playbooks/{playbook_id}/import/{document_id}/review", response_class=HTMLResponse)
async def playbook_import_review(request: Request, playbook_id: int, document_id: int, db: DBSession = Depends(get_db)):
    user = _require_user(request, db)
    playbook = _get_owned_playbook(db, user, playbook_id)
    source_document = _get_owned_source_document(db, playbook, document_id)

    clauses = []
    for clause_type in pa.CLAUSE_TYPES:
        position = pa.get_position_for_display(db, playbook.id, clause_type)
        if position is None:
            continue
        current_fields = {f.field_name: f for f in position.fields if f.superseded_by_field_id is None}
        from_this_doc = {
            name: f for name, f in current_fields.items()
            if f.evidence_document_id == source_document.id
        }
        if not from_this_doc:
            continue  # nothing from this document landed on this clause type

        established = [
            {"field_name": name, "label": pa.FIELD_LABELS[clause_type].get(name, name),
             "value": _display_value(f.value_json), "excerpt": f.evidence_excerpt, "source": f.source}
            for name, f in from_this_doc.items() if f.status == "ESTABLISHED"
        ]
        # AI-only bucket: real evidence exists, but it's qualitative or an
        # unverified quantitative claim — never silently folded into
        # "established," always a distinct, visually separate section
        # (Phase 3 task item 9).
        proposed_interpretation = [
            {"field_name": name, "label": pa.FIELD_LABELS[clause_type].get(name, name), "excerpt": f.evidence_excerpt}
            for name, f in from_this_doc.items() if f.status == "REQUIRES_LAWYER_INTERPRETATION"
        ]
        conflicting = [
            {"field_name": name, "label": pa.FIELD_LABELS[clause_type].get(name, name), "excerpt": f.evidence_excerpt}
            for name, f in from_this_doc.items() if f.status == "CONFLICTING"
        ]
        missing_labels = []
        try:
            pa.validate_position_for_activation(position)
        except pa.PolicyActivationError as exc:
            missing_labels = pa.missing_field_labels(clause_type, exc.missing_fields)

        clauses.append({
            "clause_type": clause_type, "label": pa.CLAUSE_TYPE_LABELS[clause_type],
            "position": position, "established": established,
            "proposed_interpretation": proposed_interpretation,
            "conflicting": conflicting, "needs_input": missing_labels,
        })

    return templates.TemplateResponse("playbook_import_review.html", {
        "request": request, "user": user, "playbook": playbook, "source_document": source_document,
        "clauses": clauses, "current_year": datetime.now().year,
    })


# ---------------------------------------------------------------------------
# Phase 3 — AI-Assisted Prose Playbook Import
# ---------------------------------------------------------------------------

@router.get("/playbooks/{playbook_id}/ai-import", response_class=HTMLResponse)
async def playbook_ai_import_page(request: Request, playbook_id: int, db: DBSession = Depends(get_db)):
    user = _require_user(request, db)
    playbook = _get_owned_playbook(db, user, playbook_id)
    return templates.TemplateResponse("playbook_ai_import.html", {
        "request": request, "user": user, "playbook": playbook,
        "ai_import_enabled": pai.is_ai_import_enabled(),
        "error": None, "current_year": datetime.now().year,
    })


@router.post("/playbooks/{playbook_id}/ai-import", response_class=HTMLResponse)
async def playbook_ai_import_submit(
    request: Request, playbook_id: int,
    file: UploadFile = File(...),
    consent: str = Form(""),
    db: DBSession = Depends(get_db), _csrf: None = Depends(csrf_protect),
):
    user = _require_user(request, db)
    playbook = _get_owned_playbook(db, user, playbook_id)

    def _error(message: str):
        return templates.TemplateResponse("playbook_ai_import.html", {
            "request": request, "user": user, "playbook": playbook,
            "ai_import_enabled": pai.is_ai_import_enabled(),
            "error": message, "current_year": datetime.now().year,
        }, status_code=403 if not pai.is_ai_import_enabled() else 400)

    # Server-side enforcement — checked here regardless of what the
    # submitted form contains. A disabled server never reaches the
    # consent check, the upload, or any document content.
    if not pai.is_ai_import_enabled():
        return _error("AI-assisted import is disabled for this server. Ask an administrator to enable it, or use deterministic template import instead.")
    if consent != "on":
        return _error("You must confirm the disclosure below before uploading — this document's content will be sent to the configured AI provider.")

    if not file.filename:
        return _error("Please choose a file to upload.")
    ext = os.path.splitext(file.filename.lower())[1]
    if ext not in ALLOWED_IMPORT_EXTENSIONS:
        return _error("Unsupported file type. Upload a .txt, .pdf, or .docx file.")

    file_bytes = await file.read()
    try:
        extracted_text = _extract_text_from_file(file_bytes, file.filename)
    except Exception:
        return _error("Could not read this file. Make sure it's a valid .txt, .pdf, or .docx document.")

    source_document = PlaybookSourceDocument(
        playbook_id=playbook.id, uploaded_by_user_id=user.id, document_type="LEGAL_PLAYBOOK",
        original_filename=upload_security.sanitize_filename(file.filename), extracted_text=extracted_text,
        use_as_deviation_baseline=False, use_for_policy_extraction=True,
    )
    db.add(source_document)
    db.flush()

    try:
        positions, cost_report = pai.import_ai_playbook(db, playbook, source_document, user, consent=True)
    except (pai.AIImportDisabledError, pai.AIImportConsentRequiredError) as exc:
        db.rollback()
        return _error(str(exc))

    db.commit()
    # Cost/operational metadata only -- never raw playbook text (task item 13).
    audit_log.record_event(
        db, "playbook_ai_import_completed", request=request, actor_user_id=user.id,
        target_type="playbook_source_document", target_id=source_document.id, success=True,
        metadata={"playbook_id": playbook.id, "clause_types_proposed": list(positions.keys()), **cost_report.to_metadata()},
    )
    db.commit()

    return RedirectResponse(url=f"/playbooks/{playbook.id}/import/{source_document.id}/review", status_code=302)
