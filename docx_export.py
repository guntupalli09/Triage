"""
Redlined .docx generation for the review workflow's "Generate Negotiation
Package" export — native Word Track Changes (w:ins / w:del), with a Word
comment on every accepted/edited redline explaining the deterministic rule
behind it.

python-docx has no public API for track changes or comments — both are
built here via direct OOXML element construction (docx.oxml.OxmlElement)
and, for comments specifically, direct manipulation of the saved .docx as a
zip package (python-docx can't add a new document part like word/comments.xml
on its own). This can't be validated against real Microsoft Word in this
environment; it has been checked for well-formedness, round-tripped through
python-docx's own parser, and independently validated with mammoth (a
separate, unrelated OOXML parser) confirming zero warnings/errors and
correct interpretation of the insertions, deletions, and comment
relationships. That is strong evidence, not a substitute for opening it in
Word.

Track-changes authorship is the *reviewing attorney* (an accepted/edited
redline becomes their edit the moment they accept it — the same way
accepting a Grammarly or Word Copilot suggestion attributes the resulting
change to the human, not the tool). The explanatory comment is authored as
"TriageCounsel Deterministic Engine" — the system's reasoning stays fully
visible and attributable, it just doesn't masquerade as the edit itself.

Only decisions with action in ("accepted", "edited") mutate the document —
rejected/flagged/dismissed findings belong in the cover memo
(review_workflow.build_cover_memo_text), not in the redlined document body.
"""

import io
import re
import zipfile
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional, Tuple

import review_workflow
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

COMMENTS_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
COMMENTS_RELATIONSHIP_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
ENGINE_AUTHOR = "TriageCounsel Deterministic Engine"

# XML 1.0 Char production (https://www.w3.org/TR/xml/#charsets):
#   Char ::= #x9 | #xA | #xD | [#x20-#xD7FF] | [#xE000-#xFFFD] | [#x10000-#x10FFFF]
# Anything outside that is not legal XML text content — OOXML is XML, so a
# character outside this set anywhere in a run/comment either raises
# (python-docx/lxml validate on assignment, which is how this was originally
# caught) or, worse, gets embedded unchecked into a hand-built XML string
# (see _inject_comments_part, which builds word/comments.xml by string
# formatting, not through lxml's element API) and produces a .docx that
# LOOKS like it saved successfully but contains invalid OOXML a stricter
# parser than lxml could reject later. This pattern matches every character
# NOT in the legal set, so sanitizing is `pattern.sub("", text)`.
_ILLEGAL_XML_CHARS_RE = re.compile(
    "[^\u0009\u000A\u000D\u0020-\uD7FF\uE000-\uFFFD\U00010000-\U0010FFFF]"
)


def _sanitize_xml_text(text: Optional[str]) -> str:
    """Strips characters illegal in XML 1.0 text content before they reach
    python-docx or a hand-built OOXML string. Contract text pulled from
    PDFs, scanned/OCR'd exhibits, or legacy encodings occasionally carries
    stray C0 control bytes (NUL is the common case) that are extraction
    artifacts, not meaningful legal language — removing just those
    characters preserves the surrounding text's meaning while producing
    valid XML. This does not touch legitimate whitespace (tab/LF/CR) or any
    normal Unicode text, including supplementary-plane characters (emoji,
    rare CJK, etc.) or smart quotes/em dashes, all of which are valid XML
    Char values and pass through completely unchanged."""
    if not text:
        return text or ""
    return _ILLEGAL_XML_CHARS_RE.sub("", text)


def _xml_escape(text: str) -> str:
    return (
        _sanitize_xml_text(text)
        .replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        .replace('"', "&quot;").replace("'", "&apos;")
    )


def _iso_now() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def _initials_of(author: str) -> str:
    words = [w for w in (author or "").replace(".", " ").split() if w]
    if not words:
        return "TC"
    if len(words) == 1:
        return words[0][:2].upper()
    return (words[0][0] + words[-1][0]).upper()


def _text_run(text: str, delete: bool = False):
    r = OxmlElement("w:r")
    t = OxmlElement("w:delText" if delete else "w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = _sanitize_xml_text(text)
    r.append(t)
    return r


def _ins_element(author: str, date_iso: str, rev_id: int, text: str):
    el = OxmlElement("w:ins")
    el.set(qn("w:id"), str(rev_id))
    el.set(qn("w:author"), author)
    el.set(qn("w:date"), date_iso)
    el.append(_text_run(text))
    return el


def _del_element(author: str, date_iso: str, rev_id: int, text: str):
    el = OxmlElement("w:del")
    el.set(qn("w:id"), str(rev_id))
    el.set(qn("w:author"), author)
    el.set(qn("w:date"), date_iso)
    el.append(_text_run(text, delete=True))
    return el


def _comment_range_start(comment_id: int):
    el = OxmlElement("w:commentRangeStart")
    el.set(qn("w:id"), str(comment_id))
    return el


def _comment_range_end_with_reference(comment_id: int):
    """w:commentRangeEnd plus the w:r/w:commentReference run that renders
    the clickable comment marker — always emitted together since a range
    end with no reference run is a comment nobody can open."""
    end = OxmlElement("w:commentRangeEnd")
    end.set(qn("w:id"), str(comment_id))
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rstyle = OxmlElement("w:rStyle")
    rstyle.set(qn("w:val"), "CommentReference")
    rpr.append(rstyle)
    run.append(rpr)
    ref = OxmlElement("w:commentReference")
    ref.set(qn("w:id"), str(comment_id))
    run.append(ref)
    return end, run


def _flush_text(doc: Document, para, text_chunk: str):
    """Appends text_chunk to para as plain (untracked) runs, starting new
    paragraphs on blank lines and line breaks on single newlines."""
    if not text_chunk:
        return para
    # Sanitized before splitting, not after: \n/\r are themselves legal XML
    # Char values (see _sanitize_xml_text) and are left untouched, so the
    # paragraph/line-break splitting below still sees exactly the same
    # newline structure — only illegal control characters are removed.
    text_chunk = _sanitize_xml_text(text_chunk)
    for i, block in enumerate(text_chunk.split("\n\n")):
        if i > 0:
            para = doc.add_paragraph()
        lines = block.split("\n")
        for j, line in enumerate(lines):
            if j > 0:
                para.add_run().add_break()
            if line:
                para.add_run(line)
    return para


def _inject_comments_part(docx_bytes: bytes, comments: List[Dict[str, str]]) -> bytes:
    """Adds word/comments.xml + the CommentReference character style +
    the Content_Types override + the document relationship — everything
    python-docx doesn't know how to do, done by treating the saved .docx
    as the zip package it actually is."""
    if not comments:
        return docx_bytes

    comment_items = "".join(
        f'<w:comment w:id="{c["id"]}" w:author="{_xml_escape(c["author"])}" '
        f'w:initials="{_xml_escape(c.get("initials") or _initials_of(c["author"]))}" w:date="{c["date"]}">'
        f'<w:p><w:r><w:t xml:space="preserve">{_xml_escape(c["text"])}</w:t></w:r></w:p></w:comment>'
        for c in comments
    )
    comments_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f"{comment_items}</w:comments>"
    )

    zin = zipfile.ZipFile(io.BytesIO(docx_bytes), "r")
    content_types = zin.read("[Content_Types].xml").decode("utf-8")
    rels = zin.read("word/_rels/document.xml.rels").decode("utf-8")
    styles = zin.read("word/styles.xml").decode("utf-8")

    if "/word/comments.xml" not in content_types:
        content_types = content_types.replace(
            "</Types>",
            f'<Override PartName="/word/comments.xml" ContentType="{COMMENTS_CONTENT_TYPE}"/></Types>',
        )
    if COMMENTS_RELATIONSHIP_TYPE not in rels:
        rels = rels.replace(
            "</Relationships>",
            f'<Relationship Id="rIdTriageComments" Type="{COMMENTS_RELATIONSHIP_TYPE}" Target="comments.xml"/></Relationships>',
        )
    if 'w:styleId="CommentReference"' not in styles:
        comment_ref_style = (
            '<w:style w:type="character" w:styleId="CommentReference">'
            '<w:name w:val="annotation reference"/><w:basedOn w:val="DefaultParagraphFont"/>'
            '<w:uiPriority w:val="99"/><w:semiHidden/><w:unhideWhenUsed/>'
            '<w:rPr><w:sz w:val="16"/><w:szCs w:val="16"/></w:rPr></w:style>'
        )
        styles = styles.replace("</w:styles>", comment_ref_style + "</w:styles>")

    buf_out = io.BytesIO()
    with zipfile.ZipFile(buf_out, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            if item.filename == "[Content_Types].xml":
                zout.writestr(item, content_types)
            elif item.filename == "word/_rels/document.xml.rels":
                zout.writestr(item, rels)
            elif item.filename == "word/styles.xml":
                zout.writestr(item, styles)
            else:
                zout.writestr(item, zin.read(item.filename))
        zout.writestr("word/comments.xml", comments_xml)
    return buf_out.getvalue()


def _enable_track_changes(docx_bytes: bytes) -> bytes:
    """Turns Track Changes on in the saved document's settings, so further
    edits either side makes while negotiating also get tracked automatically
    — matches how a redline is normally exchanged between firms."""
    zin = zipfile.ZipFile(io.BytesIO(docx_bytes), "r")
    settings = zin.read("word/settings.xml").decode("utf-8")
    if "<w:trackChanges" not in settings:
        idx = settings.index("<w:settings")
        close = settings.index(">", idx) + 1
        settings = settings[:close] + "<w:trackChanges/>" + settings[close:]

    buf_out = io.BytesIO()
    with zipfile.ZipFile(buf_out, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            if item.filename == "word/settings.xml":
                zout.writestr(item, settings)
            else:
                zout.writestr(item, zin.read(item.filename))
    return buf_out.getvalue()


def build_redlined_docx(
    filename: str,
    contract_text: str,
    findings: List[Dict[str, Any]],
    decisions: Dict[str, Dict[str, Any]],
    author: Optional[str] = None,
) -> Tuple[bytes, List[str]]:
    """Returns (docx_bytes, skipped_rule_ids). author is the reviewing
    attorney's name/email — becomes the Track Changes author of record for
    every accepted/edited redline (accepting one makes it their edit, the
    same way accepting a Grammarly/Copilot suggestion attributes the change
    to the human), and the author of every "reviewed and declined" comment
    on a rejected finding (that comment IS the attorney's own note — unlike
    the rule-rationale comment on an accepted redline, which stays
    attributed to the engine that computed it). Falls back to a generic
    label only if the caller genuinely has nothing better (should not
    happen for an authenticated review session).

    Two kinds of span get placed in the document, in one combined
    left-to-right pass so overlap detection covers both together:
      - "redline": accepted/edited — mutates the text (w:del the original,
        w:ins the new language) and carries a rule-rationale comment.
      - "reject_note": rejected — leaves the original text untouched and
        wraps it in a comment recording that it was reviewed and declined,
        with the attorney's stated reason. Flagged/dismissed findings (no
        redline exists to accept or reject) are not annotated here — they
        stay in the cover memo only, same as before.
    """
    author = author or "TriageCounsel Reviewer"
    decisions = decisions or {}
    date_iso = _iso_now()

    items = []
    for idx, f in enumerate(findings):
        d = decisions.get(review_workflow.finding_key(idx, f["rule_id"]))
        if not d:
            continue
        start, end = f.get("start_index"), f.get("end_index")
        if start is None or end is None or start >= end or end > len(contract_text):
            continue

        if d.get("action") in ("accepted", "edited"):
            if d["action"] == "edited":
                new_text = (d.get("edited_text") or "").strip()
            else:
                new_text = (f.get("redline") or {}).get("suggested_redline", "")
            if not new_text:
                continue
            items.append({
                "kind": "redline", "rule_id": f["rule_id"], "start": start, "end": end,
                "original": contract_text[start:end], "new_text": new_text,
                "edited": d["action"] == "edited",
                "issue": (f.get("redline") or {}).get("issue", f.get("title", f["rule_id"])),
                "rationale": (f.get("redline") or {}).get("legal_rationale") or f.get("rationale", ""),
            })
        elif d.get("action") == "rejected":
            reason = (d.get("reason") or "").strip()
            if not reason:
                continue
            items.append({
                "kind": "reject_note", "rule_id": f["rule_id"], "start": start, "end": end,
                "original": contract_text[start:end], "reason": reason,
                "issue": f.get("title", f["rule_id"]),
            })

    items.sort(key=lambda a: a["start"])
    applied, skipped = [], []
    cursor_end = -1
    for a in items:
        if a["start"] < cursor_end:
            skipped.append(a["rule_id"])
            continue
        applied.append(a)
        cursor_end = a["end"]

    doc = Document()
    doc.add_heading(f"Redlined Draft — {_sanitize_xml_text(filename)}", level=1)
    note = doc.add_paragraph()
    note_run = note.add_run(
        "This redline uses native Word Track Changes. Each suggested edit carries a comment "
        "explaining the deterministic rule behind it, and each declined suggestion carries the "
        "reviewer's note — accept, reject, or edit them exactly as you would any other tracked change."
    )
    note_run.italic = True
    doc.add_paragraph()

    comments_meta: List[Dict[str, str]] = []
    para = doc.add_paragraph()
    cursor = 0
    rev_id = 1
    comment_id = 0
    for a in applied:
        para = _flush_text(doc, para, contract_text[cursor:a["start"]])

        if a["kind"] == "redline":
            para._p.append(_del_element(author, date_iso, rev_id, a["original"]))
            rev_id += 1
            para._p.append(_comment_range_start(comment_id))
            para._p.append(_ins_element(author, date_iso, rev_id, a["new_text"]))
            rev_id += 1
            end_el, ref_run = _comment_range_end_with_reference(comment_id)
            para._p.append(end_el)
            para._p.append(ref_run)

            comment_text = f"Rule {a['rule_id']} — {a['issue']}. {a['rationale']}".strip()
            if a["edited"]:
                comment_text += " (Suggested redline was edited by the reviewer before acceptance.)"
            comments_meta.append({"id": str(comment_id), "author": ENGINE_AUTHOR, "initials": "TC", "date": date_iso, "text": comment_text})
        else:  # reject_note — text is untouched, just annotated
            para._p.append(_comment_range_start(comment_id))
            para = _flush_text(doc, para, a["original"])
            end_el, ref_run = _comment_range_end_with_reference(comment_id)
            para._p.append(end_el)
            para._p.append(ref_run)

            comment_text = f"Reviewed and declined ({a['rule_id']} — {a['issue']}): {a['reason']}"
            comments_meta.append({"id": str(comment_id), "author": author, "date": date_iso, "text": comment_text})

        comment_id += 1
        cursor = a["end"]
    _flush_text(doc, para, contract_text[cursor:])

    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()
    docx_bytes = _inject_comments_part(docx_bytes, comments_meta)
    docx_bytes = _enable_track_changes(docx_bytes)
    return docx_bytes, skipped
